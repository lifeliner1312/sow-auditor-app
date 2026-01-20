"""
Document Parser Module - Enhanced with OCR and Table Extraction
"""

import PyPDF2
from docx import Document
import os
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import re

class DocumentParser:
    """Enhanced document parser with OCR and table extraction support"""

    def __init__(self, file_path):
        self.file_path = file_path
        self.file_extension = os.path.splitext(file_path)[1].lower()
        self.content = ""
        self.tables = []
        self.is_scanned = False

    def parse(self):
        """Parse PDF or DOCX file with OCR fallback"""
        if self.file_extension == '.pdf':
            return self._parse_pdf_enhanced()
        elif self.file_extension in ['.docx', '.doc']:
            return self._parse_docx()
        else:
            raise ValueError(f"Unsupported file format: {self.file_extension}")

    def _parse_pdf_enhanced(self):
        """Enhanced PDF parsing with OCR for scanned documents"""
        try:
            # Standard text extraction with pdfplumber
            text_content = []
            tables_content = []

            with pdfplumber.open(self.file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_content.append(f"\n--- Page {page_num} ---\n{page_text}")

                    # Extract tables
                    page_tables = page.extract_tables()
                    if page_tables:
                        for table_num, table in enumerate(page_tables, 1):
                            self.tables.append({
                                'page': page_num,
                                'table_num': table_num,
                                'data': table
                            })
                            table_text = self._format_table_as_text(table)
                            tables_content.append(f"\n[TABLE {table_num} on Page {page_num}]\n{table_text}")

            combined_text = "\n".join(text_content + tables_content)

            # Check if scanned document
            word_count = len(combined_text.split())
            if word_count < 50:
                print("⚠️ Low text content detected. Attempting OCR...")
                self.is_scanned = True
                ocr_text = self._perform_ocr()
                if ocr_text:
                    combined_text = ocr_text

            self.content = combined_text
            return self.content

        except Exception as e:
            print(f"pdfplumber failed, trying PyPDF2: {str(e)}")
            return self._parse_pdf_fallback()

    def _parse_pdf_fallback(self):
        """Fallback PDF parsing using PyPDF2"""
        try:
            with open(self.file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = []
                for page in pdf_reader.pages:
                    text.append(page.extract_text())
                self.content = "\n".join(text)
                return self.content
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")

    def _perform_ocr(self):
        """Perform OCR on scanned PDF pages"""
        try:
            print("Converting PDF to images for OCR...")
            images = convert_from_path(self.file_path, dpi=300)

            ocr_text = []
            for page_num, image in enumerate(images, 1):
                print(f"Processing page {page_num}/{len(images)} with OCR...")
                page_text = pytesseract.image_to_string(image, lang='eng')
                ocr_text.append(f"\n--- Page {page_num} (OCR) ---\n{page_text}")

            return "\n".join(ocr_text)

        except Exception as e:
            print(f"OCR failed: {str(e)}")
            print("⚠️ Please ensure Tesseract OCR is installed")
            return ""

    def _format_table_as_text(self, table):
        """Format table data as readable text"""
        if not table:
            return ""

        formatted_rows = []
        for row in table:
            cleaned_row = [str(cell).strip() if cell else "" for cell in row]
            formatted_rows.append(" | ".join(cleaned_row))

        return "\n".join(formatted_rows)

    def _parse_docx(self):
        """Enhanced DOCX parsing with table support"""
        try:
            doc = Document(self.file_path)
            text = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Extract tables
            for table_num, table in enumerate(doc.tables, 1):
                table_text = f"\n[TABLE {table_num}]\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    table_text += row_text + "\n"
                text.append(table_text)

                # Store table data
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text.strip() for cell in row.cells])
                self.tables.append({
                    'table_num': table_num,
                    'data': table_data
                })

            self.content = "\n".join(text)
            return self.content

        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")

    def get_metadata(self):
        """Return file metadata"""
        file_size = os.path.getsize(self.file_path) / (1024 * 1024)

        return {
            'filename': os.path.basename(self.file_path),
            'size_mb': round(file_size, 2),
            'format': self.file_extension,
            'word_count': len(self.content.split()),
            'is_scanned': self.is_scanned,
            'tables_found': len(self.tables),
            'page_count': self._get_page_count()
        }

    def _get_page_count(self):
        """Get page count from PDF"""
        if self.file_extension == '.pdf':
            try:
                with open(self.file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    return len(pdf_reader.pages)
            except:
                return 0
        return 1

    def extract_tables(self):
        """Return extracted tables"""
        return self.tables
